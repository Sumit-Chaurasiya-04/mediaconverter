# utils/document_utils.py
# All document processing functions
# Handles PDF, Word, Excel, CSV, PowerPoint

import os
import logging
import subprocess
import shutil
from pathlib import Path
from typing import Optional, List, Dict

from utils.file_utils import generate_output_path

logger = logging.getLogger(__name__)

# PDF libraries
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.utils import ImageReader
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PIL import Image as PILImage
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# Check for LibreOffice (for DOC/PPT conversion)
LIBREOFFICE_AVAILABLE = (
    shutil.which("libreoffice") is not None or 
    shutil.which("soffice") is not None
)

def get_libreoffice_command():
    """Get the LibreOffice command (differs by OS)."""
    if shutil.which("libreoffice"):
        return "libreoffice"
    elif shutil.which("soffice"):
        return "soffice"
    return None


# ============================================================
# PDF OPERATIONS
# ============================================================

def merge_pdfs(input_paths: List[str]) -> Optional[str]:
    """
    Merge multiple PDF files into one PDF.
    
    Parameters:
        input_paths: List of PDF file paths to merge
    
    Returns:
        Path to merged PDF file
    """
    if not PYPDF2_AVAILABLE:
        logger.error("PyPDF2 not available")
        return None
    
    try:
        merger = PyPDF2.PdfMerger()
        
        for pdf_path in input_paths:
            merger.append(pdf_path)
        
        output_path = generate_output_path(input_paths[0], "pdf", "_merged")
        
        with open(output_path, "wb") as f:
            merger.write(f)
        
        merger.close()
        return output_path
        
    except Exception as e:
        logger.error(f"PDF merge error: {e}")
        return None


def split_pdf(input_path: str, start_page: int, end_page: int) -> Optional[str]:
    """
    Extract specific pages from a PDF.
    
    Parameters:
        input_path: Path to original PDF
        start_page: First page to include (1-indexed)
        end_page: Last page to include (1-indexed)
    
    Example:
        split_pdf("document.pdf", 1, 5)
        # Extracts pages 1 through 5
    """
    if not PYPDF2_AVAILABLE:
        return None
    
    try:
        reader = PyPDF2.PdfReader(input_path)
        writer = PyPDF2.PdfWriter()
        
        total_pages = len(reader.pages)
        
        # Convert to 0-indexed and clamp to valid range
        start_idx = max(0, start_page - 1)
        end_idx = min(total_pages - 1, end_page - 1)
        
        for page_num in range(start_idx, end_idx + 1):
            writer.add_page(reader.pages[page_num])
        
        output_path = generate_output_path(input_path, "pdf", f"_pages{start_page}-{end_page}")
        
        with open(output_path, "wb") as f:
            writer.write(f)
        
        return output_path
        
    except Exception as e:
        logger.error(f"PDF split error: {e}")
        return None


def compress_pdf(input_path: str) -> Optional[str]:
    """
    Compress a PDF file to reduce its size.
    Uses Ghostscript if available, otherwise PyPDF2.
    
    Ghostscript provides better compression but needs to be installed separately.
    """
    # Try Ghostscript first (better compression)
    gs_command = shutil.which("gs") or shutil.which("gswin64c") or shutil.which("gswin32c")
    
    if gs_command:
        try:
            output_path = generate_output_path(input_path, "pdf", "_compressed")
            
            command = [
                gs_command,
                "-sDEVICE=pdfwrite",
                "-dCompatibilityLevel=1.4",
                "-dPDFSETTINGS=/ebook",    # /screen=lowest, /ebook=medium, /printer=high
                "-dNOPAUSE",
                "-dQUIET",
                "-dBATCH",
                f"-sOutputFile={output_path}",
                input_path
            ]
            
            result = subprocess.run(command, capture_output=True, text=True)
            
            if result.returncode == 0 and os.path.exists(output_path):
                return output_path
        except Exception as e:
            logger.warning(f"Ghostscript compression failed: {e}, trying PyPDF2")
    
    # Fallback: PyPDF2 compression (less effective but always works)
    if PYPDF2_AVAILABLE:
        try:
            reader = PyPDF2.PdfReader(input_path)
            writer = PyPDF2.PdfWriter()
            
            for page in reader.pages:
                page.compress_content_streams()
                writer.add_page(page)
            
            output_path = generate_output_path(input_path, "pdf", "_compressed")
            
            with open(output_path, "wb") as f:
                writer.write(f)
            
            return output_path
            
        except Exception as e:
            logger.error(f"PDF compression error: {e}")
    
    return None


def get_pdf_page_count(input_path: str) -> int:
    """Get the total number of pages in a PDF."""
    if not PYPDF2_AVAILABLE:
        return 0
    
    try:
        reader = PyPDF2.PdfReader(input_path)
        return len(reader.pages)
    except:
        return 0


def pdf_to_images(input_path: str, dpi: int = 150) -> List[str]:
    """
    Convert each page of a PDF to a separate image file.
    
    Parameters:
        input_path: Path to PDF
        dpi: Resolution for output images (72=low, 150=medium, 300=high)
    
    Returns:
        List of image file paths (one per page)
    """
    output_paths = []
    
    try:
        # Try using pdf2image library first
        try:
            from pdf2image import convert_from_path
            
            images = convert_from_path(input_path, dpi=dpi)
            
            base_path = generate_output_path(input_path, "jpg")
            base_name = Path(base_path).stem
            base_dir = Path(base_path).parent
            
            for i, image in enumerate(images):
                page_path = os.path.join(str(base_dir), f"{base_name}_page{i+1}.jpg")
                image.save(page_path, "JPEG", quality=90)
                output_paths.append(page_path)
            
            return output_paths
            
        except ImportError:
            logger.warning("pdf2image not available, trying alternative method")
        
        # Fallback: use pdfplumber to render pages
        if PDFPLUMBER_AVAILABLE:
            with pdfplumber.open(input_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    img = page.to_image(resolution=dpi)
                    
                    base_path = generate_output_path(input_path, "png")
                    base_name = Path(base_path).stem
                    base_dir = Path(base_path).parent
                    
                    page_path = os.path.join(str(base_dir), f"{base_name}_page{i+1}.png")
                    img.save(page_path)
                    output_paths.append(page_path)
        
        return output_paths
        
    except Exception as e:
        logger.error(f"PDF to image error: {e}")
        return output_paths


# ============================================================
# OCR TEXT EXTRACTION
# ============================================================

def extract_text_from_pdf(input_path: str) -> Optional[str]:
    """
    Extract all text from a PDF file.
    Works best on text-based PDFs (not scanned images).
    
    Returns:
        Extracted text as a string, or None if failed
    """
    extracted_text = ""
    
    # Try pdfplumber first (better text extraction)
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(input_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        extracted_text += f"\n--- Page {page_num + 1} ---\n{text}\n"
            
            if extracted_text.strip():
                return extracted_text
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")
    
    # Fallback to PyPDF2
    if PYPDF2_AVAILABLE:
        try:
            reader = PyPDF2.PdfReader(input_path)
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    extracted_text += f"\n--- Page {page_num + 1} ---\n{text}\n"
            
            if extracted_text.strip():
                return extracted_text
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {e}")
    
    return extracted_text if extracted_text else None


def ocr_extract_text(input_path: str, language: str = "eng") -> Optional[str]:
    """
    Extract text from an image using OCR (Optical Character Recognition).
    This can read text from photos, scanned documents, screenshots, etc.
    
    Parameters:
        input_path: Path to image file (PNG, JPG, etc.)
        language: Tesseract language code ("eng" = English, "deu" = German, etc.)
    
    Returns:
        Extracted text as string, or None if Tesseract not installed
    
    Note: Requires Tesseract OCR to be installed on your computer!
          Download from: https://tesseract-ocr.github.io/tessdoc/Installation.html
    """
    if not TESSERACT_AVAILABLE:
        logger.error("pytesseract not available")
        return None
    
    try:
        img = PILImage.open(input_path)
        
        # Preprocess image for better OCR accuracy
        # Convert to grayscale - OCR works better on B&W
        if img.mode != "L":
            img = img.convert("L")
        
        # Extract text
        text = pytesseract.image_to_string(img, lang=language)
        
        return text.strip() if text.strip() else "No text found in image"
        
    except Exception as e:
        logger.error(f"OCR error: {e}")
        return None


# ============================================================
# EXCEL / CSV OPERATIONS
# ============================================================

def csv_to_excel(input_path: str) -> Optional[str]:
    """Convert a CSV file to Excel format (.xlsx)."""
    if not PANDAS_AVAILABLE:
        logger.error("Pandas not available")
        return None
    
    try:
        output_path = generate_output_path(input_path, "xlsx")
        
        # Read CSV file
        df = pd.read_csv(input_path, encoding='utf-8-sig')
        
        # Save as Excel
        df.to_excel(output_path, index=False, engine='openpyxl')
        
        return output_path
        
    except Exception as e:
        logger.error(f"CSV to Excel error: {e}")
        return None


def excel_to_csv(input_path: str, sheet_name: Optional[str] = None) -> Optional[str]:
    """
    Convert an Excel file to CSV format.
    
    Parameters:
        input_path: Path to Excel file (.xlsx or .xls)
        sheet_name: Specific sheet to convert (None = first sheet)
    """
    if not PANDAS_AVAILABLE:
        logger.error("Pandas not available")
        return None
    
    try:
        output_path = generate_output_path(input_path, "csv")
        
        # Read Excel file
        if sheet_name:
            df = pd.read_excel(input_path, sheet_name=sheet_name, engine='openpyxl')
        else:
            df = pd.read_excel(input_path, engine='openpyxl')
        
        # Save as CSV
        df.to_csv(output_path, index=False, encoding='utf-8-sig')
        
        return output_path
        
    except Exception as e:
        logger.error(f"Excel to CSV error: {e}")
        return None


def get_excel_sheets(input_path: str) -> List[str]:
    """Get list of sheet names in an Excel file."""
    if not PANDAS_AVAILABLE:
        return []
    
    try:
        xl = pd.ExcelFile(input_path, engine='openpyxl')
        return xl.sheet_names
    except:
        return []


# ============================================================
# LIBREOFFICE CONVERSIONS (DOC, PPT, etc.)
# ============================================================

def convert_with_libreoffice(input_path: str, output_format: str) -> Optional[str]:
    """
    Convert documents using LibreOffice command line.
    LibreOffice can convert between many formats including:
    - DOC/DOCX to PDF
    - PPT/PPTX to PDF
    - ODT to DOCX
    - And many more!
    
    Parameters:
        input_path: Path to input document
        output_format: Target format ("pdf", "docx", "odt", etc.)
    
    Note: LibreOffice must be installed!
          Windows: Download from https://www.libreoffice.org/
    """
    lo_command = get_libreoffice_command()
    
    if not lo_command:
        logger.error("LibreOffice not installed")
        return None
    
    try:
        import tempfile
        output_dir = tempfile.gettempdir()
        
        command = [
            lo_command,
            "--headless",              # Run without GUI
            "--convert-to", output_format,
            "--outdir", output_dir,
            input_path
        ]
        
        result = subprocess.run(command, capture_output=True, text=True, timeout=120)
        
        if result.returncode == 0:
            # LibreOffice saves to the same directory with new extension
            input_stem = Path(input_path).stem
            expected_output = os.path.join(output_dir, f"{input_stem}.{output_format}")
            
            if os.path.exists(expected_output):
                # Move to our output directory
                final_path = generate_output_path(input_path, output_format)
                shutil.move(expected_output, final_path)
                return final_path
        
        logger.error(f"LibreOffice conversion failed: {result.stderr}")
        return None
        
    except Exception as e:
        logger.error(f"LibreOffice error: {e}")
        return None


def docx_to_pdf(input_path: str) -> Optional[str]:
    """Convert a Word document (.docx) to PDF."""
    return convert_with_libreoffice(input_path, "pdf")


def pdf_to_docx_simple(input_path: str) -> Optional[str]:
    """
    Convert PDF to Word document.
    Extracts text and creates a new DOCX file.
    
    Note: Complex formatting may not be preserved perfectly.
    """
    if not (PDFPLUMBER_AVAILABLE and DOCX_AVAILABLE):
        return None
    
    try:
        # Extract text from PDF
        text_content = extract_text_from_pdf(input_path)
        
        if not text_content:
            logger.error("No text extracted from PDF")
            return None
        
        # Create Word document
        doc = Document()
        
        # Add title
        doc.add_heading("Converted from PDF", level=1)
        
        # Add extracted text
        # Split by page breaks and add as paragraphs
        pages = text_content.split("--- Page")
        
        for page_content in pages:
            if page_content.strip():
                lines = page_content.strip().split("\n")
                for line in lines:
                    if line.strip():
                        doc.add_paragraph(line.strip())
        
        output_path = generate_output_path(input_path, "docx")
        doc.save(output_path)
        
        return output_path
        
    except Exception as e:
        logger.error(f"PDF to DOCX error: {e}")
        return None


def get_document_info(input_path: str) -> dict:
    """Get information about a document file."""
    ext = Path(input_path).suffix.lower().lstrip(".")
    info = {}
    
    try:
        if ext == "pdf" and PYPDF2_AVAILABLE:
            reader = PyPDF2.PdfReader(input_path)
            info["pages"] = len(reader.pages)
            info["format"] = "PDF"
            
        elif ext in ["xlsx", "xls"] and PANDAS_AVAILABLE:
            xl = pd.ExcelFile(input_path)
            info["sheets"] = len(xl.sheet_names)
            info["sheet_names"] = xl.sheet_names
            info["format"] = "Excel"
            
        elif ext == "csv" and PANDAS_AVAILABLE:
            df = pd.read_csv(input_path)
            info["rows"] = len(df)
            info["columns"] = len(df.columns)
            info["format"] = "CSV"
            
        elif ext == "docx" and DOCX_AVAILABLE:
            doc = Document(input_path)
            info["paragraphs"] = len(doc.paragraphs)
            info["format"] = "Word Document"
            
    except Exception as e:
        logger.error(f"Error getting document info: {e}")
    
    return info